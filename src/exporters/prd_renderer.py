from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Dict, List

from docx import Document  # type: ignore[import]
from docx.enum.style import WD_STYLE_TYPE  # type: ignore[import]
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]


def load_prd(path: Path) -> Dict:
    """加载PRD JSON文件，并统一格式（兼容两种结构）"""
    prd = json.loads(path.read_text(encoding="utf-8"))
    
    # 兼容两种结构：
    # 1. 扁平结构：顶层有sections
    # 2. 嵌套结构：outputs.sections
    if "outputs" in prd and "sections" not in prd:
        # 如果是嵌套结构，扁平化处理
        outputs = prd.get("outputs", {})
        prd["sections"] = outputs.get("sections", [])
        # 如果metadata在outputs中，也提取出来
        if "metadata" not in prd:
            prd["metadata"] = prd.get("metadata", {})
    
    return prd


def render_markdown(prd: Dict, language: str = "auto") -> str:
    lines: List[str] = []
    metadata = prd.get("metadata", {})
    lines.append(f"# {metadata.get('domain', '产品需求文档')}")
    lines.append("")
    lines.append(f"- PRD ID: {metadata.get('prd_id', 'N/A')}")
    lines.append(f"- 生成时间: {metadata.get('generated_at', 'N/A')}")
    lines.append("")

    sections = prd.get("sections", [])
    for section in sections:
        section_id = section.get("section_id", "未命名章节")
        lines.append(f"## {section_id}")
        lines.append("")
        content = section.get("content", {})
        zh_text = (content.get("zh-CN") or "").strip()
        en_text = (content.get("en-US") or "").strip()

        if language in ("auto", "zh") and zh_text:
            lines.append(zh_text.strip())
            lines.append("")
        if language in ("auto", "en") and en_text:
            lines.append("> " + en_text.strip())
            lines.append("")

        for table in section.get("tables", []):
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            header_line = " | ".join(_select_text(h, language) for h in headers)
            lines.append(header_line)
            lines.append(" | ".join("---" for _ in headers))
            for row in rows:
                cells = []
                for cell in row:
                    cells.append(_cell_text(cell, language))
                lines.append(" | ".join(cells))
            lines.append("")

        if language in ("auto", "zh"):
            for figure in section.get("figures", []):
                fig_lang = figure.get("language", "zh")
                if language == "zh" and fig_lang != "zh":
                    continue
                caption = figure.get("caption", {}).get("zh-CN") or figure.get("caption", {}).get("en-US")
                image_path = figure.get("image_path") or figure.get("path")
                lines.append(f"![{caption}]({image_path})")
                lines.append("")
        if language == "en":
            for figure in section.get("figures", []):
                if figure.get("language") != "en":
                    continue
                caption = figure.get("caption", {}).get("en-US") or figure.get("caption", {}).get("zh-CN")
                image_path = figure.get("image_path") or figure.get("path")
                lines.append(f"![{caption}]({image_path})")
                lines.append("")

    return "\n".join(lines)


def _ensure_style(document: Document, style_name: str, base_style: str) -> None:
    if style_name not in document.styles:
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles[base_style]


def _setup_document_styles(document: Document) -> None:
    """设置文档样式：字体、间距、页边距等"""
    # 设置页面边距（单位：厘米，转换为英寸）
    # 使用更宽的页边距，提升阅读体验
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(2.5)      # 2.5cm ≈ 0.98 inch
        section.bottom_margin = Inches(2.5)  # 2.5cm
        section.left_margin = Inches(1.5)   # 3.8cm ≈ 1.5 inch (更宽的左边距)
        section.right_margin = Inches(1.5)  # 3.8cm (更宽的右边距)

    # 设置中文字体
    try:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Microsoft YaHei"
        normal_style.font.size = Pt(12)
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.5
    except Exception:
        pass

    # 设置标题样式
    try:
        for level in range(1, 5):
            heading_style = document.styles[f"Heading {level}"]
            heading_style.font.name = "Microsoft YaHei"
            heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if level == 1:
                heading_style.font.size = Pt(18)
                heading_style.font.bold = True
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
            elif level == 2:
                heading_style.font.size = Pt(16)
                heading_style.font.bold = True
                heading_style.paragraph_format.space_before = Pt(10)
                heading_style.paragraph_format.space_after = Pt(6)
            elif level == 3:
                heading_style.font.size = Pt(14)
                heading_style.font.bold = True
                heading_style.paragraph_format.space_before = Pt(8)
                heading_style.paragraph_format.space_after = Pt(4)
    except Exception:
        pass


def render_docx(prd: Dict, output_path: Path, language: str = "auto") -> None:
    document = Document()
    _setup_document_styles(document)
    
    _ensure_style(document, "Heading1CN", "Heading 1")
    _ensure_style(document, "Heading2CN", "Heading 2")

    metadata = prd.get("metadata", {})
    
    # ========== 设计专业的文档开头 ==========
    
    # 1. 添加顶部空白（增加视觉呼吸感）
    for _ in range(2):
        document.add_paragraph()
    
    # 2. 文档主标题（大号、加粗、居中）
    title_text = metadata.get("domain", "产品需求文档")
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text.upper())
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)  # 深蓝色
    
    # 3. 副标题（Product Requirement Document）
    subtitle_para = document.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Product Requirement Document")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # 4. 添加分隔线（装饰性）
    document.add_paragraph()
    separator_para = document.add_paragraph()
    separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator_run = separator_para.add_run("─" * 50)
    separator_run.font.size = Pt(10)
    separator_run.font.color.rgb = RGBColor(200, 200, 200)
    
    # 5. 元数据信息框（更专业的展示）
    document.add_paragraph()
    meta_table = document.add_table(rows=2, cols=2)
    meta_table.style = "Light List Accent 1"  # 使用更简洁的表格样式
    
    # 设置表格宽度和列宽
    meta_table.autofit = False
    meta_table.columns[0].width = Inches(1.2)
    meta_table.columns[1].width = Inches(4.3)
    
    # PRD ID
    prd_id_cell = meta_table.rows[0].cells[0]
    prd_id_cell.paragraphs[0].clear()
    prd_id_label = prd_id_cell.paragraphs[0].add_run("PRD ID")
    prd_id_label.bold = True
    prd_id_label.font.size = Pt(11)
    prd_id_label.font.color.rgb = RGBColor(31, 78, 121)  # 与标题颜色一致
    
    prd_id_value = metadata.get('prd_id', 'N/A')
    if len(prd_id_value) > 45:
        prd_id_value = prd_id_value[:42] + "..."
    value_cell = meta_table.rows[0].cells[1]
    value_cell.paragraphs[0].clear()
    value_cell.paragraphs[0].add_run(prd_id_value).font.size = Pt(10)
    
    # 生成时间
    time_cell = meta_table.rows[1].cells[0]
    time_cell.paragraphs[0].clear()
    time_label = time_cell.paragraphs[0].add_run("生成时间")
    time_label.bold = True
    time_label.font.size = Pt(11)
    time_label.font.color.rgb = RGBColor(31, 78, 121)
    
    generated_at = metadata.get('generated_at', 'N/A')
    # 格式化时间显示
    if generated_at != 'N/A' and 'T' in generated_at:
        try:
            from datetime import datetime
            dt_str = generated_at.replace('Z', '+00:00')
            dt = datetime.fromisoformat(dt_str)
            generated_at = dt.strftime('%Y年%m月%d日 %H:%M:%S')
        except Exception:
            pass
    
    time_value_cell = meta_table.rows[1].cells[1]
    time_value_cell.paragraphs[0].clear()
    time_value_cell.paragraphs[0].add_run(generated_at).font.size = Pt(10)
    
    # 6. 添加底部空白和分隔线（准备进入正文）
    for _ in range(2):
        document.add_paragraph()
    
    # 添加正文分隔线
    doc_separator_para = document.add_paragraph()
    doc_separator_run = doc_separator_para.add_run("═" * 60)
    doc_separator_run.font.size = Pt(8)
    doc_separator_run.font.color.rgb = RGBColor(180, 180, 180)
    
    document.add_paragraph()

    sections = prd.get("sections", [])
    for section in sections:
        section_id = section.get("section_id", "未命名章节")
        document.add_heading(section_id, level=1)
        content = section.get("content", {})
        zh_text = (content.get("zh-CN") or "").strip()
        en_text = (content.get("en-US") or "").strip()
        if language in ("auto", "zh") and zh_text:
            _write_text(document, zh_text, prefer_cn=True)
        if language in ("auto", "en") and en_text:
            _write_text(document, en_text, prefer_cn=False)

        for table in section.get("tables", []):
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if not headers or not rows:
                continue
                
            doc_table = document.add_table(rows=len(rows) + 1, cols=len(headers))
            doc_table.style = "Light Grid Accent 1"  # 更美观的表格样式
            
            # 设置表头
            header_cells = doc_table.rows[0].cells
            for idx, header in enumerate(headers):
                header_text = _clean_markdown(_select_text(header, language))
                header_para = header_cells[idx].paragraphs[0]
                header_para.clear()
                header_run = header_para.add_run(header_text)
                header_run.bold = True
                header_run.font.size = Pt(11)
                header_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置表格内容
            for r_idx, row in enumerate(rows, start=1):
                for c_idx, cell in enumerate(row):
                    cell_text = _clean_markdown(_cell_text(cell, language))
                    doc_table.rows[r_idx].cells[c_idx].text = cell_text
                    # 设置单元格对齐
                    doc_table.rows[r_idx].cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 表格前后添加间距
            document.add_paragraph()

        if language in ("auto", "zh"):
            for figure in section.get("figures", []):
                fig_lang = figure.get("language", "zh")
                if language == "zh" and fig_lang != "zh":
                    continue
                image_path = figure.get("image_path") or figure.get("path")
                caption = figure.get("caption", {}).get("zh-CN") or figure.get("caption", {}).get("en-US") or ""
                if image_path and Path(image_path).exists():
                    # 图片前添加间距
                    document.add_paragraph()
                    # 添加图片（居中，宽度6英寸）
                    pic_para = document.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = pic_para.add_run()
                    run.add_picture(image_path, width=Inches(6.0))
                    # 添加图片说明（居中，小字体）
                    if caption:
                        caption_para = document.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_para.add_run(caption)
                        caption_run.font.size = Pt(10)
                        caption_run.font.italic = True
                        caption_run.font.color.rgb = RGBColor(128, 128, 128)
                    document.add_paragraph()
                else:
                    if caption:
                        document.add_paragraph(f"[未找到图片] {caption}")
        if language in ("auto", "en"):
            for figure in section.get("figures", []):
                if figure.get("language") != "en":
                    continue
                image_path = figure.get("image_path") or figure.get("path")
                caption = figure.get("caption", {}).get("en-US") or figure.get("caption", {}).get("zh-CN") or ""
                if image_path and Path(image_path).exists():
                    # 图片前添加间距
                    document.add_paragraph()
                    # 添加图片（居中，宽度6英寸）
                    pic_para = document.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = pic_para.add_run()
                    run.add_picture(image_path, width=Inches(6.0))
                    # 添加图片说明（居中，小字体）
                    if caption:
                        caption_para = document.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_para.add_run(caption)
                        caption_run.font.size = Pt(10)
                        caption_run.font.italic = True
                        caption_run.font.color.rgb = RGBColor(128, 128, 128)
                    document.add_paragraph()
                else:
                    if caption:
                        document.add_paragraph(f"[Image not found] {caption}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def _clean_markdown(text: str) -> str:
    """清理Markdown格式字符"""
    if not text:
        return ""
    
    # 移除标题标记
    text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
    
    # 移除加粗/斜体标记（保留文本）
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # **text**
    text = re.sub(r"\*([^*]+)\*", r"\1", text)      # *text*
    text = re.sub(r"__([^_]+)__", r"\1", text)      # __text__
    text = re.sub(r"_([^_]+)_", r"\1", text)        # _text_
    
    # 移除删除线
    text = re.sub(r"~~([^~]+)~~", r"\1", text)
    
    # 移除行内代码标记
    text = re.sub(r"`([^`]+)`", r"\1", text)
    
    # 移除引用标记（>）
    text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)
    
    # 移除分隔线（--- 或 ***）
    text = re.sub(r"^[-*]{3,}\s*$", "", text, flags=re.MULTILINE)
    
    # 移除多余的空白行（最多保留一个空行）
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    # 清理行首行尾空白
    lines = [line.rstrip() for line in text.split("\n")]
    text = "\n".join(lines)
    
    return text.strip()


def _parse_rich_text(line: str, paragraph) -> None:
    """解析富文本格式并添加到段落"""
    # 处理加粗文本 **text** 或 __text__
    parts = re.split(r"(\*\*[^*]+\*\*|__[^_]+__)", line)
    
    for part in parts:
        if not part:
            continue
        # 检查是否为加粗标记
        if re.match(r"\*\*[^*]+\*\*|__[^_]+__", part):
            # 移除标记并添加加粗文本
            text = part.replace("**", "").replace("__", "")
            run = paragraph.add_run(text)
            run.bold = True
        else:
            # 普通文本
            paragraph.add_run(part)


def _write_text(document: Document, text: str, *, prefer_cn: bool) -> None:
    """
    将自由文本以更规范的段落/列表形式写入 DOCX：
    - 清理Markdown格式字符
    - 解析标题层级（##、###等）
    - 处理列表、引用块等格式
    - 优化段落间距
    """
    import json as _json

    if not text or not text.strip():
        return

    # 先清理Markdown格式
    cleaned_text = text.strip()
    
    # 1) JSON 结构友好展示（仅在明确为JSON时处理）
    if (cleaned_text.startswith("[") and cleaned_text.endswith("]")) or (cleaned_text.startswith("{") and cleaned_text.endswith("}")):
        try:
            obj = _json.loads(cleaned_text)
            if isinstance(obj, list):
                for item in obj:
                    para = document.add_paragraph(style="List Bullet")
                    para.add_run(str(item))
                return
            if isinstance(obj, dict):
                for k, v in obj.items():
                    p = document.add_paragraph()
                    run_k = p.add_run(f"{k}：")
                    run_k.bold = True
                    p.add_run(str(v))
                return
        except Exception:
            pass

    # 2) 按行处理
    lines = [ln.rstrip() for ln in cleaned_text.split("\n") if ln.strip()]
    
    if not lines:
        return

    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 检查标题层级（## 或 ###）- 先处理标题标记，但保留其他格式
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # 清理标题中的Markdown格式
            heading_text = _clean_markdown(heading_text)
            document.add_heading(heading_text, level=min(level, 3))
            i += 1
            continue
        
        # 检查列表项（- * • · 或数字列表）
        list_match = re.match(r"^([-*•·]|\d+\.)\s+(.+)$", line)
        if list_match:
            bullet_text = list_match.group(2).strip()
            para = document.add_paragraph(style="List Bullet")
            # 解析富文本（加粗等）但清理其他Markdown格式
            _parse_rich_text(bullet_text, para)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        
        # 检查子列表项（缩进的列表）
        sub_list_match = re.match(r"^\s{2,}([-*•·]|\d+\.)\s+(.+)$", line)
        if sub_list_match:
            bullet_text = sub_list_match.group(2).strip()
            para = document.add_paragraph(style="List Bullet 2")
            _parse_rich_text(bullet_text, para)
            para.paragraph_format.left_indent = Inches(1.0)
            para.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        
        # 检查引用块（>）
        if line.startswith("> "):
            quote_text = line[2:].strip()
            # 清理引用块中的Markdown格式
            quote_text = _clean_markdown(quote_text)
            para = document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(64, 64, 64)
            i += 1
            continue
        
        # 检查分隔线（---）
        if re.match(r"^-{3,}$", line):
            # 添加段落间距
            document.add_paragraph()
            i += 1
            continue
        
        # 普通段落 - 清理Markdown但保留加粗等格式
        cleaned_line = _clean_markdown(line)
        if cleaned_line:
            para = document.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            # 对于普通段落，直接添加文本（已清理Markdown）
            para.add_run(cleaned_line)
        
        i += 1


def _select_text(item: Dict, language: str) -> str:
    if language == "en":
        return item.get("en-US") or item.get("zh-CN") or item.get("value", "")
    if language == "zh":
        return item.get("zh-CN") or item.get("en-US") or item.get("value", "")
    return item.get("zh-CN") or item.get("en-US") or item.get("value", "")


def _cell_text(cell: Dict | str, language: str) -> str:
    if isinstance(cell, dict):
        return _select_text(cell, language)
    return str(cell)


